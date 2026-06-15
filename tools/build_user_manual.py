from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Manual_de_Usuario_DenunciaPE.docx"
LOGO = ROOT / "frontend" / "public" / "LogoDenunciaPE.png"
ICON = ROOT / "frontend" / "public" / "IconoDenunciaPE-limpio.png"

RED = "C81A1A"
DARK_RED = "8F1212"
INK = "172033"
MUTED = "64748B"
LIGHT_RED = "FEF2F2"
LIGHT_GRAY = "F1F5F9"
MID_GRAY = "CBD5E1"
WHITE = "FFFFFF"
GREEN = "047857"
AMBER = "B45309"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_image_alt_text(inline_shape, description):
    inline_shape._inline.docPr.set("descr", description)
    inline_shape._inline.docPr.set("title", description)


def set_run(run, size=11, color=INK, bold=False, italic=False, font="Aptos"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_run(run, 9, MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run(lead, bold=True)
        run = p.add_run(text[len(bold_lead):])
        set_run(run)
    else:
        run = p.add_run(text)
        set_run(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.19)
        p.paragraph_format.space_after = Pt(4)
        set_run(p.add_run(item))


def add_steps(doc, items):
    for title, detail in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.19)
        p.paragraph_format.space_after = Pt(5)
        set_run(p.add_run(f"{title}. "), bold=True)
        set_run(p.add_run(detail))


def add_callout(doc, title, text, tone="info"):
    palette = {
        "info": (LIGHT_RED, RED),
        "warning": ("FFF7ED", AMBER),
        "success": ("ECFDF5", GREEN),
        "neutral": (LIGHT_GRAY, MUTED),
    }
    fill, accent = palette[tone]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.35)
    cell = table.cell(0, 0)
    prevent_row_split(table.rows[0])
    cell.width = Inches(6.35)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 140, 180, 140, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(title), 10.5, accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    set_run(p2.add_run(text), 10.5, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for idx, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(width)
        set_cell_shading(cell, RED)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(header), 9.5, WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, (value, width) in enumerate(zip(row, widths)):
            cells[idx].width = Inches(width)
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2:
                set_cell_shading(cells[idx], "FAFAFA")
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx else WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(str(value)), 9.5, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_screen_path(doc, path):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run("Ruta en la plataforma: "), 9.5, MUTED, bold=True)
    set_run(p.add_run(path), 9.5, RED, bold=True)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, color, before, after in [
        ("Heading 1", 17, RED, 16, 8),
        ("Heading 2", 13.5, DARK_RED, 12, 6),
        ("Heading 3", 11.5, INK, 9, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if ICON.exists():
        icon_shape = header_p.add_run().add_picture(str(ICON), height=Inches(0.28))
        set_image_alt_text(icon_shape, "Icono de DenunciaPE")
    set_run(header_p.add_run("  DenunciaPE | Manual de usuario"), 9, MUTED, bold=True)

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        logo_shape = p.add_run().add_picture(str(LOGO), width=Inches(3.6))
        set_image_alt_text(logo_shape, "Logotipo de DenunciaPE")
    elif ICON.exists():
        icon_shape = p.add_run().add_picture(str(ICON), height=Inches(1.15))
        set_image_alt_text(icon_shape, "Icono de DenunciaPE")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run("MANUAL DE USUARIO"), 28, RED, bold=True, font="Aptos Display")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    set_run(p.add_run("Plataforma digital para el registro y seguimiento de denuncias"), 15, INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Guía para ciudadanos, policías, encargados de comisaría y Super Administrador"), 11, MUTED)

    doc.add_paragraph()
    add_callout(
        doc,
        "Objetivo del documento",
        "Explicar de forma práctica cómo acceder a DenunciaPE, verificar la identidad, registrar una denuncia, seleccionar el lugar del hecho, consultar su estado y utilizar el panel institucional.",
        "neutral",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(26)
    set_run(p.add_run("Versión 1.0 | Junio de 2026 | Perú"), 10, MUTED)
    doc.add_page_break()


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "Contenido", 1)
    toc = [
        "1. Acerca de DenunciaPE",
        "2. Requisitos y recomendaciones",
        "3. Primer ingreso y verificación de identidad",
        "4. Panel ciudadano",
        "5. Registro de una nueva denuncia",
        "6. Seguimiento y consulta de denuncias",
        "7. Panel institucional",
        "8. Funciones por rol",
        "9. Seguridad, privacidad y buenas prácticas",
        "10. Solución de problemas",
        "11. Glosario y estados de la denuncia",
    ]
    add_bullets(doc, toc)
    add_callout(doc, "Alcance", "Este manual describe la versión actual del prototipo web DenunciaPE. Algunas integraciones gubernamentales y validaciones biométricas avanzadas están simuladas para fines de demostración.", "warning")

    add_heading(doc, "1. Acerca de DenunciaPE", 1)
    add_body(doc, "DenunciaPE es una plataforma web que permite registrar digitalmente denuncias por robo o hurto, validar la identidad del denunciante y consultar el avance del caso sin acudir inicialmente a una comisaría.")
    add_heading(doc, "1.1 Usuarios de la plataforma", 2)
    add_table(
        doc,
        ["Usuario", "Funciones principales"],
        [
            ("Ciudadano", "Verificar su identidad, registrar denuncias, revisar expedientes y hacer seguimiento."),
            ("Policía", "Consultar denuncias disponibles, revisar expedientes y aceptar casos de su comisaría."),
            ("Encargado", "Supervisar denuncias y administrar cuentas policiales de su comisaría."),
            ("Fiscal", "Consultar únicamente los expedientes derivados a su despacho."),
            ("Super Admin", "Supervisar todas las denuncias, comisarías y cuentas institucionales."),
        ],
        [1.35, 5.0],
    )

    add_heading(doc, "2. Requisitos y recomendaciones", 1)
    add_bullets(doc, [
        "Dispositivo con navegador actualizado: computadora, tableta o teléfono.",
        "Conexión a Internet estable.",
        "DNI, fecha de nacimiento y datos de contacto.",
        "Acceso al correo electrónico registrado.",
        "Cámara o archivos de imagen para la verificación facial.",
        "Información del hecho: fecha, ubicación, objetos sustraídos y relato.",
    ])
    add_callout(doc, "Antes de empezar", "Reúne fotos, videos, comprobantes, datos de testigos y características de los objetos. La plataforma guarda el borrador, pero tener la información preparada reduce errores.", "info")

    add_heading(doc, "3. Primer ingreso y verificación de identidad", 1)
    add_heading(doc, "3.1 Iniciar el proceso ciudadano", 2)
    add_screen_path(doc, "Inicio > Empezar denuncia / Iniciar sesión")
    add_steps(doc, [
        ("Selecciona el acceso ciudadano", "Usa la opción para empezar una denuncia si todavía no tienes acceso registrado, o inicia sesión con tu DNI/correo y contraseña."),
        ("Completa tus datos", "Registra nombres, apellidos, DNI, fecha de nacimiento, correo, celular opcional y una contraseña segura."),
        ("Revisa los datos", "Comprueba que el DNI tenga ocho dígitos, el celular nueve dígitos y el correo esté escrito correctamente."),
    ])

    add_heading(doc, "3.2 Tutorial inicial", 2)
    add_body(doc, "En el primer ingreso puedes revisar un tutorial de seis pasos. Explica qué es una denuncia, la diferencia entre robo y hurto, los requisitos, las evidencias permitidas, el proceso posterior y el seguimiento.")
    add_callout(doc, "Robo y hurto", "Robo: existe violencia, amenaza o fuerza contra la persona. Hurto: el bien es sustraído sin violencia, por ejemplo sin que la víctima lo advierta.", "warning")

    add_heading(doc, "3.3 Verificación del correo", 2)
    add_steps(doc, [
        ("Solicita el código", "La plataforma envía un código de seis dígitos al correo registrado."),
        ("Revisa tu bandeja", "Busca el mensaje en la bandeja principal y en correo no deseado."),
        ("Ingresa el código", "Escríbelo en DenunciaPE y selecciona Verificar correo."),
        ("Reenvía si es necesario", "Si el código venció o no llegó, usa Reenviar código."),
    ])

    add_heading(doc, "3.4 Consentimiento y verificación facial", 2)
    add_body(doc, "Antes de usar la cámara, DenunciaPE explica el propósito de las imágenes y solicita consentimiento expreso. En la versión actual se capturan tres imágenes: frontal, lateral izquierda y lateral derecha.")
    add_steps(doc, [
        ("Acepta el consentimiento", "Confirma que autorizas el uso limitado de las imágenes para verificar tu identidad."),
        ("Prepara el entorno", "Busca buena iluminación, retira lentes oscuros y mantén el rostro visible."),
        ("Realiza las capturas", "Sigue la guía de pantalla para la foto frontal y ambos perfiles."),
        ("Finaliza", "Cuando las tres capturas se registren, accederás al panel ciudadano."),
    ])
    add_callout(doc, "Privacidad", "Las imágenes faciales no deben reutilizarse para otros fines. El sistema registra el consentimiento y protege los archivos mediante almacenamiento restringido.", "info")

    add_heading(doc, "4. Panel ciudadano", 1)
    add_body(doc, "Después de completar la verificación, el panel presenta tres acciones principales:")
    add_table(
        doc,
        ["Opción", "Uso"],
        [
            ("Nueva denuncia", "Inicia el formulario guiado de robo o hurto."),
            ("Mis denuncias", "Muestra borradores y denuncias registradas con su estado y oficina actual."),
            ("Consultar por código", "Permite consultar la ubicación administrativa de una denuncia mediante su código."),
        ],
        [1.75, 4.6],
    )

    add_heading(doc, "5. Registro de una nueva denuncia", 1)
    add_screen_path(doc, "Panel ciudadano > Nueva denuncia")
    add_body(doc, "El formulario contiene ocho pasos obligatorios y guarda el avance como borrador. El policía asistente muestra recomendaciones contextuales en cada sección.")

    add_heading(doc, "5.1 Paso 1: tipo de hecho", 2)
    add_steps(doc, [
        ("Selecciona Robo", "Cuando hubo amenaza, violencia o fuerza."),
        ("Selecciona Hurto", "Cuando el bien fue sustraído sin violencia."),
        ("Continúa", "La clasificación puede revisarse antes del envío final."),
    ])

    add_heading(doc, "5.2 Paso 2: fecha y lugar", 2)
    add_steps(doc, [
        ("Indica fecha y hora", "No se permite una fecha futura."),
        ("Selecciona el ubigeo", "Elige departamento, provincia y distrito."),
        ("Espera el centrado", "El mapa se desplazará automáticamente al distrito seleccionado."),
        ("Marca el punto exacto", "Haz clic en el mapa. Se guardarán latitud y longitud."),
        ("Revisa la referencia", "La plataforma sugerirá una calle o zona. Puedes corregirla o completarla manualmente."),
    ])
    add_callout(doc, "Cambio de distrito", "Si cambias departamento, provincia o distrito, el punto y la referencia anteriores se eliminan. Debes seleccionar nuevamente el lugar exacto.", "warning")

    add_heading(doc, "5.3 Paso 3: relato de los hechos", 2)
    add_body(doc, "Describe qué ocurrió, cómo se produjo el hecho y qué sucedió después. El relato debe tener al menos 30 caracteres y puede contener hasta 2,000.")
    add_bullets(doc, [
        "Usa un orden cronológico.",
        "Incluye acciones observables y evita suposiciones.",
        "Menciona amenazas, violencia, armas, vehículos o dirección de huida.",
        "No incluyas contraseñas, claves bancarias ni datos innecesarios.",
    ])

    add_heading(doc, "5.4 Paso 4: objetos sustraídos", 2)
    add_body(doc, "Registra al menos un objeto. Para cada uno indica nombre, marca/modelo o característica, cantidad y valor aproximado.")
    add_callout(doc, "Ejemplo", "Celular Samsung A54, color negro, una unidad, valor aproximado S/ 1,200.", "neutral")

    add_heading(doc, "5.5 Paso 5: sospechosos", 2)
    add_body(doc, "Indica si observaste a los sospechosos. Si respondes Sí, completa la descripción física, vestimenta y forma de huida. Si no los viste, selecciona No.")

    add_heading(doc, "5.6 Paso 6: testigos", 2)
    add_body(doc, "Indica si hubo testigos. Si respondes Sí, registra nombre, relación y teléfono de nueve dígitos. Informa al testigo que sus datos serán incluidos en la denuncia.")

    add_heading(doc, "5.7 Paso 7: evidencias", 2)
    add_body(doc, "Puedes adjuntar fotografías, videos, capturas, boletas u otros archivos relacionados. Agrega una descripción breve para facilitar su revisión.")
    add_bullets(doc, [
        "Adjunta únicamente archivos relevantes.",
        "Evita imágenes duplicadas o borrosas.",
        "No alteres ni edites evidencia que deba conservar su autenticidad.",
        "La evidencia es opcional para continuar, pero puede ayudar a la investigación.",
    ])

    add_heading(doc, "5.8 Paso 8: revisión y envío", 2)
    add_steps(doc, [
        ("Revisa el resumen", "Comprueba tipo, fecha, ubicación, coordenadas, relato, objetos, sospechosos, testigos y evidencias."),
        ("Acepta las declaraciones", "Confirma la veracidad de la información y el tratamiento de datos."),
        ("Envía la denuncia", "El sistema validará que los campos obligatorios estén completos."),
        ("Guarda el código", "Al finalizar se muestra un código como DEN-2026-0001234."),
        ("Descarga la constancia", "Usa la opción disponible para conservar la constancia provisional."),
    ])
    add_callout(doc, "Confirmación", "No consideres registrada la denuncia hasta visualizar el mensaje Denuncia registrada y el código de seguimiento.", "success")

    add_heading(doc, "6. Seguimiento y consulta de denuncias", 1)
    add_heading(doc, "6.1 Desde Mis denuncias", 2)
    add_screen_path(doc, "Panel ciudadano > Mis denuncias > Seleccionar expediente")
    add_body(doc, "La pantalla muestra el código, tipo, distrito, estado y oficina actual. Al abrir un expediente puedes revisar el recorrido por oficinas, el relato, objetos, sospechosos y testigos.")

    add_heading(doc, "6.2 Consulta sin iniciar sesión", 2)
    add_screen_path(doc, "Inicio > Hacer seguimiento")
    add_steps(doc, [
        ("Ingresa el código", "Escribe el código completo, incluyendo guiones."),
        ("Selecciona Consultar", "El sistema mostrará únicamente información pública de seguimiento."),
        ("Revisa la oficina actual", "Verás la comisaría, la oficina y los movimientos administrativos."),
    ])
    add_callout(doc, "Protección de datos", "La consulta pública no muestra el relato, datos del denunciante, testigos, evidencias ni objetos sustraídos.", "info")

    add_heading(doc, "7. Panel institucional", 1)
    add_screen_path(doc, "Inicio institucional > Iniciar sesión")
    add_steps(doc, [
        ("Ingresa tus credenciales", "Usa el usuario, correo institucional o DNI y la contraseña asignada."),
        ("Revisa el alcance", "El panel identifica tu rol y comisaría o cobertura nacional."),
        ("Consulta el resumen", "Visualiza totales, casos recibidos, en investigación y resueltos."),
        ("Abre un expediente", "Selecciona Ver expediente para consultar datos, objetos, testigos, evidencias y trazabilidad."),
    ])
    add_callout(doc, "Acceso restringido", "Las cuentas institucionales son creadas por el Super Administrador o por el encargado autorizado. No compartas credenciales.", "warning")

    add_heading(doc, "8. Funciones por rol", 1)
    add_heading(doc, "8.1 Policía", 2)
    add_steps(doc, [
        ("Abre Denuncias disponibles", "Revisa los casos sin responsable policial dentro de tu comisaría."),
        ("Consulta el expediente", "Verifica los datos antes de asumir responsabilidad."),
        ("Selecciona Aceptar denuncia", "El caso cambia a Asignada y queda registrado en la trazabilidad."),
        ("Usa Denuncias a mi cargo", "Consulta los expedientes que aceptaste."),
    ])

    add_heading(doc, "8.2 Encargado de comisaría", 2)
    add_bullets(doc, [
        "Consulta las denuncias asignadas a su comisaría.",
        "Revisa el resumen operativo de la sede.",
        "Crea cuentas para policías de su propia comisaría.",
        "Consulta el personal registrado dentro de su alcance.",
    ])

    add_heading(doc, "8.3 Fiscal", 2)
    add_body(doc, "El fiscal accede únicamente a los expedientes derivados formalmente a su despacho. No puede consultar denuncias de otras comisarías ni crear cuentas.")

    add_heading(doc, "8.4 Super Administrador", 2)
    add_bullets(doc, [
        "Supervisa denuncias a nivel nacional.",
        "Consulta todas las comisarías y sus indicadores.",
        "Abre una comisaría para revisar sus denuncias y personal.",
        "Registra nuevas comisarías.",
        "Crea cuentas de policías, encargados y fiscales.",
    ])
    add_heading(doc, "Registrar una comisaría", 3)
    add_screen_path(doc, "Panel institucional > Comisarías > Nueva comisaría")
    add_steps(doc, [
        ("Completa el nombre", "Usa el nombre institucional de la sede."),
        ("Registra la ubicación", "Indica departamento, provincia, distrito y dirección."),
        ("Agrega una referencia", "Este campo es opcional."),
        ("Confirma el registro", "La sede aparecerá en el listado y podrá recibir personal y denuncias."),
    ])

    add_heading(doc, "Registrar personal institucional", 3)
    add_screen_path(doc, "Panel institucional > Personal > Nueva cuenta")
    add_body(doc, "Completa usuario, nombres, apellidos, DNI, correo institucional y contraseña temporal. El Super Administrador selecciona rol y comisaría; el encargado solo puede crear policías para su propia sede.")

    add_heading(doc, "9. Seguridad, privacidad y buenas prácticas", 1)
    add_bullets(doc, [
        "Usa una contraseña exclusiva y no la compartas.",
        "Cierra sesión al terminar, especialmente en equipos compartidos.",
        "Verifica la dirección web antes de ingresar datos.",
        "No envíes códigos de verificación ni contraseñas por mensajería.",
        "Adjunta únicamente información necesaria para la denuncia.",
        "Evita descargar expedientes en dispositivos no autorizados.",
        "Reporta accesos o cambios que no reconozcas.",
    ])
    add_callout(doc, "Datos sensibles", "La información de identidad, imágenes faciales, testimonios y evidencias debe tratarse conforme a la Ley N.° 29733 y a las políticas institucionales aplicables.", "info")

    add_heading(doc, "10. Solución de problemas", 1)
    add_table(
        doc,
        ["Problema", "Qué hacer"],
        [
            ("No llega el código", "Revisa correo no deseado, confirma el correo y usa Reenviar código."),
            ("La cámara no abre", "Autoriza la cámara en el navegador, usa HTTPS o selecciona una imagen permitida."),
            ("El mapa no carga", "Comprueba Internet. Aun si falla la dirección automática, mueve el mapa y marca el punto."),
            ("El mapa apunta a otro lugar", "Confirma departamento, provincia y distrito; luego selecciona el punto exacto."),
            ("No puedo continuar", "Revisa el mensaje rojo y completa todos los campos obligatorios del paso."),
            ("Perdí el código", "Inicia sesión y abre Mis denuncias para recuperarlo."),
            ("Sesión vencida", "Vuelve a iniciar sesión. El borrador local puede restaurarse automáticamente."),
            ("Denuncia no encontrada", "Verifica el código completo y los guiones. Si persiste, contacta a soporte institucional."),
        ],
        [2.0, 4.35],
    )

    add_heading(doc, "11. Glosario y estados de la denuncia", 1)
    add_table(
        doc,
        ["Estado", "Significado"],
        [
            ("Recibida", "La denuncia fue registrada correctamente."),
            ("Pendiente de verificación de identidad", "La identidad requiere validación adicional o revisión asistida."),
            ("En revisión", "Personal autorizado está verificando la información."),
            ("Asignada", "Un policía asumió la responsabilidad del caso."),
            ("Información adicional requerida", "Se necesita información complementaria del denunciante."),
            ("En investigación", "El caso se encuentra en trabajo operativo."),
            ("Resuelta", "El flujo institucional registró la conclusión del caso."),
            ("Observada", "Se detectó una inconsistencia que debe revisarse."),
        ],
        [2.25, 4.1],
    )
    add_heading(doc, "Glosario", 2)
    add_bullets(doc, [
        "Código de seguimiento: identificador único entregado al registrar la denuncia.",
        "Constancia provisional: documento generado por el MVP; no sustituye documentos oficiales que la PNP determine.",
        "Trazabilidad: historial de oficinas, responsables y movimientos del expediente.",
        "Ubigeo: selección de departamento, provincia y distrito.",
        "Geocodificación: proceso que centra el mapa según el ubigeo y convierte coordenadas en una referencia aproximada.",
        "Borrador: denuncia iniciada que todavía no fue enviada.",
    ])

    add_heading(doc, "Guía rápida", 1)
    add_table(
        doc,
        ["Ciudadano", "Personal institucional"],
        [
            ("1. Registrarse o iniciar sesión.", "1. Iniciar sesión institucional."),
            ("2. Verificar correo y rostro.", "2. Revisar resumen y alcance."),
            ("3. Crear una nueva denuncia.", "3. Abrir la bandeja correspondiente."),
            ("4. Completar los ocho pasos.", "4. Revisar el expediente."),
            ("5. Guardar el código.", "5. Aceptar o gestionar según el rol."),
            ("6. Consultar Mis denuncias.", "6. Mantener la trazabilidad del caso."),
        ],
        [3.175, 3.175],
    )
    add_callout(doc, "Fin del manual", "DenunciaPE guía al ciudadano antes, durante y después del registro, y prepara información estructurada para la atención institucional.", "success")

    doc.core_properties.title = "Manual de Usuario - DenunciaPE"
    doc.core_properties.subject = "Uso de la plataforma de denuncias digitales DenunciaPE"
    doc.core_properties.author = "Equipo DenunciaPE"
    doc.core_properties.keywords = "DenunciaPE, manual, denuncias, robo, hurto, Perú"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
